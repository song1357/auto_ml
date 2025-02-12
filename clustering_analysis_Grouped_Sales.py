import pandas as pd
import numpy as np
from kmodes.kmodes import KModes
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io
import os
from sklearn.preprocessing import MinMaxScaler
import time
from docx.oxml import parse_xml
from sklearn.metrics import silhouette_score

plt.rcParams['font.sans-serif'] = ['SimHei']  # Use SimHei font for Chinese characters
plt.rcParams['axes.unicode_minus'] = False     # Fix minus sign display

def create_clustering_report():
    # 创建Word文档
    doc = Document()
    doc.add_heading('调味品使用聚类分析报告', 0)
    
    # 读取数据
    excel_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "D:/GoalYeah_Tasks/2-10/聚类的基础数据(3).xlsx")
    df_sku = pd.read_excel(excel_file, sheet_name="店01")

    # 数据预处理（静默进行，不输出到报告）
    X = df_sku.iloc[:, 3:].copy()
    X = X.replace([' ', '', 'nan', 'NaN', 'NULL'], 0).infer_objects(copy=False)
    X = X.fillna(0)
    for column in X.columns:
        X[column] = pd.to_numeric(X[column], errors='coerce').fillna(0)
    
    # 聚类分析
    doc.add_heading('1. 聚类数量评估', level=1)
    
    # 计算不同K值的评估指标
    K = range(2, 11)
    costs = []  # KModes的cost值（类似于WCSS）
    silhouette_scores = []
    
    for k in K:
        # 执行KModes聚类
        km = KModes(n_clusters=k, init='Huang', n_init=5, random_state=42)
        clusters = km.fit_predict(X)
        costs.append(km.cost_)
        
        # 计算轮廓系数
        silhouette_scores.append(silhouette_score(X, clusters, metric='hamming'))
    
    # 绘制评估图
    plt.figure(figsize=(12, 5))
    
    # 肘部法则图（使用KModes的cost值）
    plt.subplot(1, 2, 1)
    plt.plot(K, costs, 'bx-')
    plt.xlabel('k值 (聚类数量)')
    plt.ylabel('组内平方和')
    plt.title('肘部法则图')
    
    # 轮廓系数图
    plt.subplot(1, 2, 2)
    plt.plot(K, silhouette_scores, 'rx-')
    plt.xlabel('k值 (聚类数量)')
    plt.ylabel('轮廓系数')
    plt.title('轮廓系数分析')
    plt.tight_layout()
    
    # 保存评估图到Word
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    doc.add_picture(buf, width=Inches(7))
    plt.close()
    
    # 添加评估图解释
    doc.add_heading('1.1 评估图解释', level=2)
    
    # 创建评估图解释表格
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格标题
    cells = table.rows[0].cells
    cells[0].text = '肘部法则图解释'
    cells[1].text = '轮廓系数解释'
    
    # 设置表头样式
    for cell in cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = 1
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
        tcPr.append(tcVAlign)
    
    # 添加解释内容
    elbow_content = """1. 肘部法则通过观察组内平方和（WCSS）的变化来评估聚类效果。
2. WCSS表示组内样本与其质心之间的距离平方和，值越小表示聚类越紧密。
3. 当增加聚类数量时，WCSS会持续下降。
4. 在某个K值处，WCSS的下降速率会显著减缓，形成"肘部"。
5. 这个"肘部"位置通常被认为是较好的聚类数量选择。"""

    silhouette_content = """1. 轮廓系数评估聚类的内聚度和分离度，取值范围为[-1, 1]。
2. 值越接近1，表示样本与自己所在的组更相似，与其他组更不相似。
3. 值越接近0，表示样本位于组的边界。
4. 值越接近-1，表示样本可能被分配到了错误的组。
5. 一般认为轮廓系数大于0.5表示聚类效果较好。"""

    cells = table.rows[1].cells
    cells[0].text = elbow_content
    cells[1].text = silhouette_content
    
    # 设置内容单元格样式
    for cell in cells:
        cell.paragraphs[0].alignment = 1
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
        tcPr.append(tcVAlign)
    
    # 对K=6到10进行详细分析
    doc.add_heading('2. 不同K值的聚类分析', level=1)
    
    for k in range(6, 11):
        doc.add_heading(f'K={k}的聚类分析', level=2)
        
        # 执行聚类
        km = KModes(n_clusters=k, init='Huang', n_init=5, random_state=42)
        cluster_labels = km.fit_predict(X)
        
        # 添加聚类标签
        df_analysis = df_sku.copy()
        df_analysis['cluster'] = cluster_labels
        
        # 分析聚类特征
        cluster_features = []
        for i in range(k):
            cluster_data = df_analysis[df_analysis['cluster'] == i].iloc[:, 3:-1]
            cluster_data = cluster_data.apply(pd.to_numeric, errors='coerce')
            mean_values = cluster_data.mean()
            cluster_features.append(mean_values)
        
        # 创建热力图
        plt.figure(figsize=(15, 8))
        sns.heatmap(pd.DataFrame(cluster_features, columns=X.columns),
                   cmap='YlOrRd',
                   annot=True,
                   fmt='.2f',
                   cbar_kws={'label': '平均值'})
        plt.title(f'K={k} 各组的调味品使用特征分析')
        plt.xlabel('调味品')
        plt.ylabel('组别')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        
        # 保存热力图
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        doc.add_picture(buf, width=Inches(7))
        plt.close()
        
        # 添加SKU占比表格标题
        title_paragraph = doc.add_paragraph('调味品使用占比（本类别中有使用的店数/本类别的总店数）')
        title_paragraph.alignment = 1  # 居中对齐
        title_paragraph.runs[0].bold = True  # 加粗标题

        # 创建SKU占比表格
        sku_table = doc.add_table(rows=k+1, cols=len(X.columns)+1)
        sku_table.style = 'Table Grid'

        # 设置表头
        header_cells = sku_table.rows[0].cells
        header_cells[0].text = '聚类'
        for idx, col in enumerate(X.columns):
            header_cells[idx+1].text = col

        # 设置表头样式
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True  # 加粗
            cell.paragraphs[0].alignment = 1  # 居中对齐

        # 填充表格内容
        for i in range(k):
            row_cells = sku_table.rows[i+1].cells
            
            # 聚类编号
            row_cells[0].text = f'聚类{i}'
            
            # 获取当前聚类的数据
            cluster_data = df_analysis[df_analysis['cluster'] == i]
            total_stores = len(cluster_data)
            
            # 计算每个调味品的使用比例
            for col_idx, col in enumerate(X.columns):
                # 计算有使用该调味品的店铺数量
                stores_using_sku = (cluster_data.iloc[:, 3:][col] > 0).sum()
                if stores_using_sku > 0:
                    percentage = (stores_using_sku / total_stores) * 100
                    row_cells[col_idx+1].text = f'{int(round(percentage))}%'
                else:
                    row_cells[col_idx+1].text = '/'

        # 设置表格样式
        for row in sku_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = 1  # 居中对齐
                # 设置单元格垂直居中
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                tcPr.append(tcVAlign)

        doc.add_paragraph('')  # 添加空行作为分隔
         # 添加标题
        doc.add_heading(f'聚类 K={k} 的菜系和餐厅档次分布', level=2)
        
        # 为每个聚类进行菜系和餐厅档次分析
        for i in range(k):
            # 添加聚类标题
            doc.add_paragraph(f'聚类{i}')
            
            # 创建分布分析表格
            distribution_table = doc.add_table(rows=1, cols=3)
            distribution_table.style = 'Table Grid'
            
            # 设置表头
            header_cells = distribution_table.rows[0].cells
            header_cells[0].text = '类型'
            header_cells[1].text = '名称'
            header_cells[2].text = '占比'

            # 设置表头样式
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True  # 加粗
                cell.paragraphs[0].alignment = 1  # 居中对齐
            
            # 获取当前聚类的数据
            cluster_data = df_analysis[df_analysis['cluster'] == i]
            
            # 获取菜系和餐厅档次列
            cuisine_type = cluster_data.iloc[:, 1]  # 第2列：菜系
            restaurant_level = cluster_data.iloc[:, 2]  # 第3列：餐厅档次
            
            # 1. 菜系分布
            cuisine_dist = cuisine_type.value_counts()
            top_cuisines = cuisine_dist.head(3)  # 只展示前3个最常见的菜系
            
            # 2. 餐厅档次分布
            level_dist = restaurant_level.value_counts()
            top_levels = level_dist  # 展示所有档次
            
            # 添加数据到表格
            for dist_type, dist_data in [
                ('菜系', top_cuisines),
                ('餐厅档次', top_levels)
            ]:
                first_row = True
                for name, value in dist_data.items():
                    row = distribution_table.add_row()
                    cells = row.cells
                    cells[0].text = str(dist_type)
                    cells[1].text = str(name)
                    
                    # 计算百分比
                    total = len(cluster_data)
                    percentage = (value / total) * 100
                    cells[2].text = f'{int(round(percentage))}%'
                    
                    # 设置单元格样式
                    for cell in cells:
                        cell.paragraphs[0].alignment = 1  # 居中对齐
                        # 设置单元格垂直居中
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                        tcPr.append(tcVAlign)
            
            doc.add_paragraph('')  # 添加空行作为分隔
        
        doc.add_paragraph('')  # 添加空行作为分隔
    
    # 保存文档
    current_dir = os.path.dirname(os.path.abspath(__file__))
    try:
        output_file = os.path.join(current_dir, '聚类分析报告_调味品使用.docx')
        doc.save(output_file)
        print(f"分析完成！报告已保存为'{output_file}'")
    except Exception as e:
        import tempfile
        temp_dir = tempfile.gettempdir()
        output_file = os.path.join(temp_dir, '聚类分析报告_调味品使用.docx')
        try:
            doc.save(output_file)
            print(f"分析完成！由于权限问题，报告已保存为'{output_file}'")
        except Exception as e:
            print(f"保存文件时出错: {str(e)}")
            # 尝试使用一个不同的临时文件名
            output_file = os.path.join(temp_dir, f'聚类分析报告_调味品使用_{int(time.time())}.docx')
            doc.save(output_file)
            print(f"分析完成！报告已保存为备用位置：'{output_file}'")

if __name__ == "__main__":
    create_clustering_report()

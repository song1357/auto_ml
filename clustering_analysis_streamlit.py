import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx2pdf import convert
import io
from sklearn.preprocessing import StandardScaler
from kmodes.kmodes import KModes
from sklearn.cluster import KMeans, MiniBatchKMeans
from sklearn.metrics import silhouette_score
import tempfile
import os
from docx.oxml.ns import qn
from functools import lru_cache
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
from docx.shared import Cm

plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

# 添加缓存装饰器来优化性能
@st.cache_data
def load_data(file):
    try:
        xls = pd.ExcelFile(file)
        sheet_names = xls.sheet_names
        all_data = {}
        for i, sheet in enumerate(sheet_names):
            all_data[sheet] = pd.read_excel(file, sheet_name=sheet)
        return all_data, sheet_names
    except Exception as e:
        st.error(f'加载数据时出错: {str(e)}')
        return None, None

@st.cache_data
def preprocess_data(df, feature_columns):
    """
    预处理数据，只保留数值型特征
    """
    # 选择指定的特征列
    X = df[feature_columns].copy()
    
    # 只保留数值型列
    numeric_cols = X.select_dtypes(include=['int64', 'float64']).columns
    if len(numeric_cols) < len(feature_columns):
        st.warning(f"以下列包含非数值数据，将被忽略: {', '.join(set(feature_columns) - set(numeric_cols))}")
    X = X[numeric_cols]
    
    # 处理缺失值
    X = X.fillna(0)
    
    return X

@st.cache_data
def get_clustering_model(model_name, n_clusters):
    if model_name == 'KModes':
        return KModes(n_clusters=n_clusters, init='Huang', n_init=5, verbose=0)
    elif model_name == 'MiniBatchKMeans':
        return MiniBatchKMeans(n_clusters=n_clusters, batch_size=1024, random_state=42)
    else:
        return KMeans(n_clusters=n_clusters, random_state=42)

@st.cache_data
def find_optimal_k(X, max_k, model_name):
    """
    寻找最优的聚类数量
    """
    K = range(2, max_k + 1)
    silhouette_scores = []
    costs = []
    
    # 标准化数据
    if model_name in ['KMeans', 'MiniBatchKMeans']:
        scaler = StandardScaler()
        X_values = scaler.fit_transform(X)
    else:
        X_values = X.copy()
    
    for k in K:
        # 根据模型名称选择不同的聚类模型
        if model_name == 'KModes':
            kmeans = KModes(n_clusters=k, init='Huang', n_init=5, verbose=0)
        elif model_name == 'MiniBatchKMeans':
            kmeans = MiniBatchKMeans(n_clusters=k, batch_size=1024, random_state=42)
        else:
            kmeans = KMeans(n_clusters=k, random_state=42)
        
        # 执行聚类
        kmeans.fit(X_values)
        
        # 计算轮廓系数
        labels = kmeans.labels_
        score = silhouette_score(X_values, labels) if len(np.unique(labels)) > 1 else 0
        silhouette_scores.append(score)
        
        # 计算组内平方和
        if hasattr(kmeans, 'inertia_'):
            costs.append(kmeans.inertia_)
        else:
            # 对于KModes，使用cost_属性的总和作为成本
            costs.append(kmeans.cost_)
    
    # 使用轮廓系数确定最优k值
    optimal_k = K[np.argmax(silhouette_scores)]
    
    return K, silhouette_scores, costs, optimal_k

def create_evaluation_plots(K, costs, silhouette_scores, optimal_k, current_k):
    """
    创建评估图，包括肘部法则图和轮廓系数图
    增加current_k参数来标记当前选择的k值
    """
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
    
    # 肘部法则图
    ax1.plot(K, costs, 'bx-')
    ax1.axvline(x=current_k, color='r', linestyle='--', alpha=0.5, label=f'当前选择: k={current_k}')
    ax1.set_xlabel('聚类数量 (k)')
    ax1.set_ylabel('组内平方和 (WCSS)')
    ax1.set_title('肘部法则评估图')
    ax1.grid(True)
    ax1.legend()

    # 轮廓系数图
    ax2.plot(K, silhouette_scores, 'rx-')
    ax2.axvline(x=current_k, color='b', linestyle='--', alpha=0.5, label=f'当前选择: k={current_k}')
    ax2.set_xlabel('聚类数量 (k)')
    ax2.set_ylabel('轮廓系数')
    ax2.set_title('轮廓系数评估图')
    ax2.grid(True)
    ax2.legend()

    plt.tight_layout()
    return fig

def create_clustering_report(df, filename, cuisine_col_idx, level_col_idx, selected_features, k_range, model_name, output_format='docx', selected_cuisines='全部', selected_level='全部'):
    # 保存原始数据用于聚类分析
    df_for_clustering = df.copy()
    
    # 根据选择的餐厅档次筛选数据用于聚类
    if selected_level != '全部':
        df_for_clustering = df[df.iloc[:, level_col_idx] == selected_level].copy()
        if len(df_for_clustering) == 0:
            raise ValueError(f"筛选后没有符合条件的数据：{selected_level}")
    
    doc = Document()
    doc.add_heading('调味品使用聚类分析报告', 0)
    doc.add_heading('1. 基本信息', level=1)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Table Grid'
    header_cells = info_table.rows[0].cells
    header_cells[0].text = '项目'
    header_cells[1].text = '内容'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = 1
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
        tcPr.append(tcVAlign)

    # 处理菜系和档次的显示格式
    if isinstance(selected_cuisines, list):
        cuisines_str = ', '.join(selected_cuisines)
    else:
        cuisines_str = selected_cuisines

    info_data = [
        ('数据文件', filename),
        ('选择的特征', ', '.join(selected_features)),
        ('聚类范围', f'K = {min(k_range)} ~ {max(k_range)}'),
        ('分析时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('选择的菜系', cuisines_str),
        ('选择的档次', selected_level)
    ]
    for label, value in info_data:
        row = info_table.add_row()
        cells = row.cells
        cells[0].text = label
        cells[1].text = str(value)
        for cell in cells:
            cell.paragraphs[0].alignment = 1
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
            tcPr.append(tcVAlign)

    # 预处理数据，只保留数值型特征
    X = preprocess_data(df_for_clustering, selected_features)
    
    doc.add_heading('2. 聚类数量评估', level=1)
    K, silhouette_scores, costs, optimal_k = find_optimal_k(X, max(k_range), model_name)
    fig = create_evaluation_plots(K, costs, silhouette_scores, optimal_k, optimal_k)
    img_path = os.path.join(tempfile.gettempdir(), 'evaluation.png')
    plt.savefig(img_path, dpi=300, bbox_inches='tight')
    plt.close()
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_before = Pt(8)
    paragraph_format.space_after = Pt(12)
    run = paragraph.add_run()
    run.add_picture(img_path, width=Inches(7))
    os.remove(img_path)

    doc.add_heading('2.1 评估图解释', level=2)
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = '肘部法则图解释'
    cells[1].text = '轮廓系数解释'
    for cell in cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = 1
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
        tcPr.append(tcVAlign)

    elbow_content = """\
    1. 肘部法则通过观察组内平方和（WCSS）的变化来评估聚类效果。
    2. WCSS表示组内样本与其质心之间的距离平方和，值越小表示聚类越紧密。
    3. 当增加聚类数量时，WCSS会持续下降。
    4. 在某个K值处，WCSS的下降速率会显著减缓，形成"肘部"。
    5. 这个"肘部"位置通常被认为是较好的聚类数量选择。"""

    silhouette_content = """\
    1. 轮廓系数评估聚类的内聚度和分离度，取值范围为[-1, 1]。
    2. 值越接近1，表示样本与自己所在的组更相似，与其他组更不相似。
    3. 值越接近0，表示样本位于组的边界。
    4. 值越接近-1，表示样本可能被分配到了错误的组。
    5. 一般认为轮廓系数大于0.5表示聚类效果较好。"""

    cells = table.rows[1].cells
    cells[0].text = elbow_content
    cells[1].text = silhouette_content
    for cell in cells:
        cell.paragraphs[0].alignment = 1
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
        tcPr.append(tcVAlign)

    doc.add_heading('3. 不同K值的聚类分析', level=1)
    cuisine_col = df.columns[cuisine_col_idx]
    level_col = df.columns[level_col_idx]

    for k in k_range:
        doc.add_heading(f'K={k}的聚类分析', level=2)
        
        # 使用预处理后的数据进行聚类
        if model_name == 'KModes':
            model = KModes(n_clusters=k, init='Huang', n_init=5, verbose=0)
        elif model_name == 'MiniBatchKMeans':
            model = MiniBatchKMeans(n_clusters=k, batch_size=1024, random_state=42)
        else:
            model = KMeans(n_clusters=k, random_state=42)
        
        clusters = model.fit_predict(X)
        df_analysis = df_for_clustering.copy()
        df_analysis['cluster'] = clusters

        # 计算每个簇的特征平均值（只使用数值型列）
        cluster_features = []
        for i in range(k):
            cluster_data = df_analysis[df_analysis['cluster'] == i][selected_features]
            # 确保所有列都是数值型
            numeric_data = cluster_data.apply(pd.to_numeric, errors='coerce')
            mean_values = numeric_data.mean()
            cluster_features.append(mean_values)

        # 绘制热力图
        plt.figure(figsize=(15, 8))
        sns.heatmap(pd.DataFrame(cluster_features, columns=selected_features),
                   cmap='YlOrRd',
                   annot=True,
                   fmt='.2f',
                   cbar_kws={'label': '平均值'})
        plt.title(f'K={k} 各组的调味品使用特征分析')
        plt.xlabel('调味品')
        plt.ylabel('组别')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()

        # 保存图片
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=300, bbox_inches='tight')
        plt.close()

        # 添加图片到文档
        paragraph = doc.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(12)
        run = paragraph.add_run()
        run.add_picture(buf, width=Inches(7))

        # 添加使用占比表格
        title_paragraph = doc.add_paragraph('调味品使用占比（本类别中有使用的店数/本类别的总店数）')
        title_paragraph.alignment = 1
        title_paragraph.runs[0].bold = True

        sku_table = doc.add_table(rows=k+1, cols=len(selected_features)+2)  # 增加一列用于店数占比
        sku_table.style = 'Table Grid'
        header_cells = sku_table.rows[0].cells
        header_cells[0].text = '聚类'
        header_cells[1].text = '店数占比'  # 新增店数占比列
        for idx, col in enumerate(selected_features):
            header_cells[idx+2].text = col

        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = 1
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
            tcPr.append(tcVAlign)

        total_stores = len(df_analysis)  # 总店数
        for i in range(k):
            row_cells = sku_table.rows[i+1].cells
            row_cells[0].text = f'聚类{i}'
            cluster_data = df_analysis[df_analysis['cluster'] == i]
            cluster_stores = len(cluster_data)
            # 添加店数占比
            store_percentage = (cluster_stores / total_stores) * 100
            row_cells[1].text = f'{int(round(store_percentage))}%'
            
            for col_idx, col in enumerate(selected_features):
                # 确保数据是数值型
                numeric_data = pd.to_numeric(cluster_data[col], errors='coerce')
                stores_using_sku = (numeric_data > 0).sum()
                if stores_using_sku > 0:
                    percentage = (stores_using_sku / cluster_stores) * 100
                    row_cells[col_idx+2].text = f'{int(round(percentage))}%'
                else:
                    row_cells[col_idx+2].text = '/'
        
        for row in sku_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].alignment = 1
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                tcPr.append(tcVAlign)

        doc.add_paragraph('')

        # 添加菜系和餐厅档次分布分析
        doc.add_heading(f'聚类 K={k} 的菜系和餐厅档次分布', level=2)
        for i in range(k):
            doc.add_paragraph(f'聚类{i}')
            distribution_table = doc.add_table(rows=1, cols=3)
            distribution_table.style = 'Table Grid'
            header_cells = distribution_table.rows[0].cells
            header_cells[0].text = '类型'
            header_cells[1].text = '名称'
            header_cells[2].text = '占比'

            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = 1
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                tcPr.append(tcVAlign)

            cluster_indices = df_analysis[df_analysis['cluster'] == i].index
            cluster_data = df.iloc[cluster_indices]
            
            # 获取菜系和餐厅档次列
            cuisine_col = cluster_data.columns[cuisine_col_idx]
            level_col = cluster_data.columns[level_col_idx]
            
            # 1. 菜系分布
            cuisine_dist = cluster_data[cuisine_col].value_counts()
            top_cuisines = cuisine_dist.head(3)
            
            # 2. 餐厅档次分布
            level_dist = cluster_data[level_col].value_counts()
            
            # 添加数据到表格
            for dist_type, dist_data in [('菜系', top_cuisines), ('餐厅档次', level_dist)]:
                if dist_type == '餐厅档次':
                    # 如果选择了特定档次，只显示该档次的数据
                    if selected_level != '全部':
                        if selected_level in level_dist:
                            row = distribution_table.add_row()
                            cells = row.cells
                            cells[0].text = dist_type
                            cells[1].text = selected_level
                            # 使用当前聚类的总数作为计算百分比的基数
                            total = len(cluster_data)  # 使用所有档次的总数
                            percentage = (level_dist[selected_level] / total) * 100
                            cells[2].text = f'{int(round(percentage))}%'
                            
                            # 设置单元格格式
                            for cell in cells:
                                cell.paragraphs[0].alignment = 1
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                                tcPr.append(tcVAlign)
                    else:
                        # 如果没有选择特定档次，显示所有档次数据
                        for name, value in dist_data.items():
                            row = distribution_table.add_row()
                            cells = row.cells
                            cells[0].text = dist_type
                            cells[1].text = str(name)
                            total = len(cluster_data)
                            percentage = (value / total) * 100
                            cells[2].text = f'{int(round(percentage))}%'
                            
                            # 设置单元格格式
                            for cell in cells:
                                cell.paragraphs[0].alignment = 1
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                                tcPr.append(tcVAlign)
                else:
                    # 菜系分布的处理保持不变
                    for name, value in dist_data.items():
                        row = distribution_table.add_row()
                        cells = row.cells
                        cells[0].text = dist_type
                        cells[1].text = str(name)
                        total = len(cluster_data)
                        percentage = (value / total) * 100
                        cells[2].text = f'{int(round(percentage))}%'
                        
                        # 设置单元格格式
                        for cell in cells:
                            cell.paragraphs[0].alignment = 1
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcVAlign = parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>')
                            tcPr.append(tcVAlign)
            
            doc.add_paragraph('')

        doc.add_paragraph('')

    return save_report(doc, output_format)

def apply_document_style(doc):
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        style.font.name = '微软雅黑'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.font.size = Pt(16 - (i-1)*2)
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(10)

def save_report(doc, output_format='docx'):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            docx_path = tmp_docx.name
        doc.save(docx_path)
        if output_format == 'pdf':
            try:
                pythoncom.CoInitialize()
                pdf_path = docx_path.replace('.docx', '.pdf')
                convert(docx_path, pdf_path)
                with open(pdf_path, 'rb') as pdf_file:
                    encoded_pdf = pdf_file.read()
                os.remove(pdf_path)
                os.remove(docx_path)
                pythoncom.CoUninitialize()
                return encoded_pdf, 'pdf'
            except Exception as e:
                pythoncom.CoUninitialize()
                st.warning(f'PDF 转换失败，将使用 Word 格式: {str(e)}')
                with open(docx_path, 'rb') as docx_file:
                    encoded_docx = docx_file.read()
                os.remove(docx_path)
                return encoded_docx, 'docx'
        else:
            with open(docx_path, 'rb') as docx_file:
                encoded_docx = docx_file.read()
            os.remove(docx_path)
            return encoded_docx, 'docx'
    except Exception as e:
        st.error(f'保存报告时发生错误: {str(e)}')
        raise

def get_cuisine_types(df):
    """获取菜系类型列表"""
    # 获取第二列的列名（B列）
    cuisine_column = df.columns[1]  # 索引1对应第二列（B列）
    st.info(f"使用列 '{cuisine_column}' 作为菜系分类")
    cuisine_types = sorted(df[cuisine_column].unique().tolist())
    return cuisine_types

def get_restaurant_levels():
    return ['全部', '中高档', '大众']

def main():
    st.set_page_config(page_title="调味品使用聚类分析", layout="wide")
    st.title("调味品-聚类分析")

    # 在main函数开始时添加cuisine_column变量
    if 'cuisine_column' not in st.session_state:
        st.session_state.cuisine_column = None

    uploaded_file = st.file_uploader("上传Excel文件", type=['xlsx', 'xls'])
    
    if uploaded_file is not None:
        try:
            all_data, sheet_names = load_data(uploaded_file)
            if all_data is None or sheet_names is None:
                return

            # 选择要分析的sheet
            if 'selected_sheet' not in st.session_state:
                st.session_state.selected_sheet = sheet_names[0]
            
            selected_sheet = st.selectbox(
                "选择要分析的sheet",
                options=sheet_names,
                key='sheet_selector',
                index=sheet_names.index(st.session_state.selected_sheet)
            )

            if selected_sheet != st.session_state.selected_sheet:
                st.session_state.selected_sheet = selected_sheet
                # 重置其他相关的session_state
                if f'cuisine_select_{selected_sheet}' in st.session_state:
                    del st.session_state[f'cuisine_select_{selected_sheet}']
                if f'level_select_{selected_sheet}' in st.session_state:
                    del st.session_state[f'level_select_{selected_sheet}']

            df = all_data[selected_sheet]
            
            # 更新cuisine_column（获取第二列的列名）
            st.session_state.cuisine_column = df.columns[1]
            
            cuisine_types = get_cuisine_types(df)
            restaurant_levels = get_restaurant_levels()

            # 餐厅档次选择
            selected_level = st.selectbox('选择餐厅档次', restaurant_levels, index=0, key='selected_level')
            
            # 菜系多选控件，默认全选
            def update_cuisines():
                # 获取当前选中的菜系并保存到对应工作表的状态中
                current_selected = st.session_state[f'cuisine_select_{selected_sheet}']
                st.session_state[f'selected_cuisines_{selected_sheet}'] = current_selected
            
            # 初始化当前工作表的菜系选择状态
            if f'selected_cuisines_{selected_sheet}' not in st.session_state:
                st.session_state[f'selected_cuisines_{selected_sheet}'] = cuisine_types
            
            # 使用工作表特定的状态作为默认值
            selected_cuisines = st.multiselect(
                label='选择菜系',
                options=cuisine_types,
                default=st.session_state[f'selected_cuisines_{selected_sheet}'],
                on_change=update_cuisines,
                key=f'cuisine_select_{selected_sheet}'
            )
            
            # 获取D列到P列的所有特征（调味品）
            feature_columns = list(df.iloc[:, 3:16].columns)  # D列是index 3，P列是index 15
            
            # 初始化当前工作表的调味品选择状态
            if f'selected_seasonings_{selected_sheet}' not in st.session_state:
                st.session_state[f'selected_seasonings_{selected_sheet}'] = feature_columns
            
            def update_seasonings():
                # 获取当前选中的调味品并保存到对应工作表的状态中
                current_selected = st.session_state[f'seasoning_select_{selected_sheet}']
                st.session_state[f'selected_seasonings_{selected_sheet}'] = current_selected
            
            # 使用container来控制多选控件的高度
            with st.container():
                selected_features = st.multiselect(
                    '选择调味品',
                    options=feature_columns,
                    default=st.session_state[f'selected_seasonings_{selected_sheet}'],
                    on_change=update_seasonings,
                    key=f'seasoning_select_{selected_sheet}'
                )
            
            # 创建两列布局：左侧为控件区域，右侧为数据预览
            col_controls, col_preview = st.columns([1, 2])
            
            # 右侧列：数据预览
            with col_preview:
                if 'selected_sheet' in st.session_state:
                    # 根据选择过滤数据
                    filtered_df = df.copy()
                    
                    # 菜系过滤：只保留选中的菜系数据
                    if st.session_state.get(f'cuisine_select_{st.session_state["selected_sheet"]}'):
                        filtered_df = filtered_df[filtered_df[st.session_state.cuisine_column].isin(st.session_state[f'cuisine_select_{st.session_state["selected_sheet"]}'])]
                    else:  # 如果没有选中任何菜系，显示空数据框
                        filtered_df = filtered_df.head(0)
                    
                    # 餐厅档次过滤
                    if st.session_state.get('selected_level') != '全部':
                        filtered_df = filtered_df[filtered_df['餐厅档次L'] == st.session_state['selected_level']]
            
                    # 获取特征列
                    feature_columns = list(df.iloc[:, 3:16].columns)  # D列是index 3，P列是index 15
                    
                    # 显示数据预览
                    # 使用选中的调味品列更新预览列
                    preview_columns = [df.columns[0], st.session_state.cuisine_column, '餐厅档次L']
                    if st.session_state.get(f'selected_seasonings_{selected_sheet}'):
                        preview_columns.extend(st.session_state[f'selected_seasonings_{selected_sheet}'])
                    
                    preview_table = st.empty()
                    with preview_table.container():
                        st.dataframe(
                            filtered_df[preview_columns],
                            use_container_width=True,
                            height=570,
                            key=f"preview_{st.session_state['selected_sheet']}_{hash(str(filtered_df.values.tobytes()))}"
                        )
                        st.write(f"当前显示数据量：{len(filtered_df)} 条")
            
                    if len(filtered_df) == 0:
                        st.warning("没有符合条件的数据")
                        st.stop()
            
            # 特征选择和聚类模型控件（放在两列布局下方）
            if 'selected_sheet' in st.session_state:
                selected_sheet = st.session_state['selected_sheet']
                if f'selected_seasonings_{selected_sheet}' in st.session_state and len(st.session_state[f'selected_seasonings_{selected_sheet}']) > 0:
                    # 聚类模型选择
                    model_name = st.selectbox(
                        "选择聚类模型",
                        ['KMeans', 'MiniBatchKMeans', 'KModes'],
                        help="KMeans: 适用于数值型数据，基于欧氏距离\nMiniBatchKMeans: KMeans的优化版本，适用于大规模数据\nKModes: 适用于分类数据，基于汉明距离"
                    )
                    
                    # 添加开始聚类分析按钮
                    if st.button('开始聚类分析'):
                        with st.spinner('正在进行聚类分析...'):
                            # 预处理数据
                            X = preprocess_data(all_data[st.session_state['selected_sheet']], st.session_state[f'selected_seasonings_{selected_sheet}'])
                            
                            # 聚类分析
                            K, silhouette_scores, costs, optimal_k = find_optimal_k(X, 10, model_name)
                            
                            # 保存聚类结果到session state
                            st.session_state['clustering_results'] = {
                                'K': list(K),  # 转换为列表以便序列化
                                'silhouette_scores': silhouette_scores,
                                'costs': costs,
                                'optimal_k': optimal_k,
                                'X': X,
                                'model_name': model_name
                            }
                            
                    # 如果已经完成聚类分析，显示报告生成选项
                    if 'clustering_results' in st.session_state:       
                        # 选择聚类数量
                        k_value = st.slider('选择聚类数量', 
                                          min_value=2, 
                                          max_value=10, 
                                          value=st.session_state['clustering_results']['optimal_k'])
                        
                        st.write("- **肘部法则图**：在拐点处的聚类数量通常是较好的选择")
                        st.write("- **轮廓系数图**：轮廓系数越大（越接近1）表示聚类效果越好")
                        
                        results = st.session_state['clustering_results']
                        fig = create_evaluation_plots(
                            results['K'],
                            results['costs'],
                            results['silhouette_scores'],
                            results['optimal_k'],
                            k_value  # 传入当前选择的k值
                        )
                        st.pyplot(fig)

                        # 选择聚类数量范围
                        col1, col2 = st.columns(2)
                        with col1:
                            k_min = st.number_input('最小聚类数', 
                                                  min_value=2, 
                                                  max_value=10, 
                                                  value=2)
                        with col2:
                            k_max = st.number_input('最大聚类数', 
                                                  min_value=2, 
                                                  max_value=10, 
                                                  value=max(k_min, k_value))
                        
                        if k_max < k_min:
                            st.warning('最大聚类数不能小于最小聚类数，已自动调整为最小聚类数。')
                            k_max = k_min
                        # 选择报告格式
                        output_format = st.radio(
                            "选择报告格式",
                            ['docx', 'pdf'],
                            format_func=lambda x: 'Word文档' if x == 'docx' else 'PDF文档',
                            horizontal=True
                        )
                        
                        if st.button('生成分析报告'):
                            try:
                                with st.spinner('正在生成报告...'):
                                    # 获取当前选择的餐厅档次
                                    current_level = st.session_state.get('selected_level', '全部')
                                    
                                    # 调用create_clustering_report时传递选择的档次
                                    report_content, file_extension = create_clustering_report(
                                        all_data[st.session_state['selected_sheet']], 
                                        f"{uploaded_file.name}_{st.session_state['selected_sheet']}",
                                        1, 
                                        2, 
                                        st.session_state[f'selected_seasonings_{selected_sheet}'],
                                        range(k_min, k_max + 1),
                                        st.session_state['clustering_results']['model_name'],
                                        output_format,
                                        st.session_state[f'selected_cuisines_{st.session_state["selected_sheet"]}'], 
                                        current_level  # 传递选择的档次
                                    )
                                    extension = '.pdf' if file_extension == 'pdf' else '.docx'
                                    filename = f'聚类分析报告_{datetime.now().strftime("%Y%m%d_%H%M%S")}{extension}'
                                    st.download_button(
                                        label="下载报告",
                                        data=report_content,
                                        file_name=filename,
                                        mime=('application/pdf' if file_extension == 'pdf' 
                                             else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                                    )
                                    st.success('报告生成成功！')
                            except Exception as e:
                                st.error(f'生成报告时发生错误: {str(e)}')
                                raise
        except Exception as e:
            st.error(f"处理数据时发生错误: {str(e)}")
            return

if __name__ == '__main__':
    main()